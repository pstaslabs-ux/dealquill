"""
CRE Deal Analyzer Dashboard — Powered by Claude
================================================
Upload deal documents → Claude extracts data → Live investment dashboard
showing Returns, 50% Rule, and 30-year projections.
"""

import io
import json
import math

from dq_utils import (
    get_history_file,
    load_history as _load_history_util,
    fmt_d,
    fmt_p,
    calc_payment,
    populate_sidebar_from_data,
    infer_units_from_type,
)
import os
import re
import uuid
from datetime import datetime

import anthropic
import pandas as pd
import plotly.graph_objects as go
import streamlit as st

# ── Config ────────────────────────────────────────────────────────────────────

try:
    API_KEY = st.secrets.get("ANTHROPIC_API_KEY") or os.environ.get("ANTHROPIC_API_KEY", "")
    HUD_TOKEN = st.secrets.get("HUD_TOKEN") or os.environ.get("HUD_TOKEN", "")
except Exception:
    API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
    HUD_TOKEN = os.environ.get("HUD_TOKEN", "")
MODEL = "claude-opus-4-6"

def load_history():
    return _load_history_util()


def save_to_history(data):
    prop  = data.get("property", {})
    fin   = data.get("financing", {})
    addr  = prop.get("address") or "Unknown Address"
    price = prop.get("purchase_price")
    label = addr + (f" — ${price:,.0f}" if price else "")
    entry = {
        "id":        str(uuid.uuid4()),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "label":     label,
        "address":   addr,
        "price":     price,
        "data":      data,
    }
    history = load_history()
    # Avoid duplicate consecutive saves of the same address+price
    if history and history[0].get("label") == label:
        history[0] = entry
    else:
        history.insert(0, entry)
    history = history[:100]  # keep last 100
    try:
        with open(get_history_file(), "w") as f:
            json.dump(history, f)
    except Exception:
        pass

# ── Page setup ────────────────────────────────────────────────────────────────

st.set_page_config(page_title="DealQuill", layout="wide", page_icon="🏢")

st.markdown("""
<style>
/* Force white background and black text everywhere */
html, body, [class*="css"], .stApp, .stApp > div,
[data-testid="stAppViewContainer"], [data-testid="stHeader"],
[data-testid="stSidebar"], [data-testid="stSidebar"] > div,
[data-testid="stMain"], .main, .block-container {
    background-color: #ffffff !important;
    color: #000000 !important;
}

/* All text elements */
p, span, div, label, h1, h2, h3, h4, h5, h6,
.stMarkdown, .stText, [data-testid="stMarkdownContainer"],
[data-testid="stCaptionContainer"] {
    color: #000000 !important;
}

/* Inputs and widgets */
input, textarea, select,
[data-testid="stTextInput"] input,
[data-testid="stNumberInput"] input,
.stSelectbox select {
    background-color: #ffffff !important;
    color: #000000 !important;
    border: 1px solid #cccccc !important;
}

/* Sidebar */
section[data-testid="stSidebar"] * {
    color: #000000 !important;
    background-color: #f8f8f8 !important;
}
section[data-testid="stSidebar"] input {
    background-color: #ffffff !important;
}

/* Metric cards */
[data-testid="metric-container"] {
    background: #ffffff !important;
    border: 1px solid #dddddd !important;
    border-radius: 10px;
    padding: 14px 18px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
}
[data-testid="metric-container"] label,
[data-testid="metric-container"] p,
[data-testid="metric-container"] span {
    color: #000000 !important;
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
    font-size: 26px !important;
    font-weight: 700 !important;
    color: #000000 !important;
}

/* Section titles */
.section-title {
    font-size: 17px; font-weight: 700; color: #000000 !important;
    border-bottom: 2px solid #000000; padding-bottom: 5px;
    margin: 22px 0 14px 0;
}

/* Dataframe */
[data-testid="stDataFrame"] * {
    color: #000000 !important;
    background-color: #ffffff !important;
}

/* Missing info box */
.missing-box {
    background: #fffbe6 !important; border: 1px solid #ccaa00 !important;
    border-radius: 8px; padding: 14px; margin-top: 16px;
    font-size: 14px; color: #000000 !important;
}

/* Expander */
[data-testid="stExpander"] *, details *, summary * {
    color: #000000 !important;
    background-color: #ffffff !important;
}

/* Hide GitHub source button in header */
header a[href*="github"],
[data-testid="stHeader"] a[href*="github"],
[data-testid="stToolbar"] a[href*="github"] {
    display: none !important;
}

/* Buttons */
.stButton button {
    color: #2F5496 !important;
    background-color: #2F5496 !important;
    border: none !important;
    outline: none !important;
    box-shadow: none !important;
}

/* DealQuill fixed header brand */
.dealquill-brand {
    position: fixed;
    top: 0;
    left: 60px;
    height: 48px;
    display: flex;
    align-items: center;
    gap: 0px;
    z-index: 999999;
    pointer-events: none;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="dealquill-brand">
    <span style="font-size:20px;font-weight:800;color:#2F5496;letter-spacing:-0.5px;">Deal</span><span style="font-size:20px;font-weight:800;color:#111;letter-spacing:-0.5px;">Quill</span>
</div>
""", unsafe_allow_html=True)


# ── Auth gate ─────────────────────────────────────────────────────────────────

def _get_users():
    """Return {username: password} from secrets or env fallback."""
    try:
        users = st.secrets.get("users")
        if users:
            return dict(users)
    except Exception:
        pass
    # Fallback: single-user mode via APP_PASSWORD
    pw = os.environ.get("APP_PASSWORD", "")
    try:
        pw = pw or st.secrets.get("APP_PASSWORD") or ""
    except Exception:
        pass
    return {"admin": pw} if pw else {}

def _check_login():
    if st.session_state.get("authenticated"):
        return True
    users = _get_users()
    if not users:
        st.session_state["username"] = "default"
        return True

    col_l, col_c, col_r = st.columns([1, 2, 1])
    with col_c:
        st.markdown("### Sign in to DealQuill")
        username = st.text_input("Username", key="_login_user").strip().lower()
        pw_input = st.text_input("Password", type="password", key="_login_pw")
        if st.button("Sign In", type="primary", use_container_width=True):
            if username in users and pw_input == users[username]:
                st.session_state["authenticated"] = True
                st.session_state["username"] = username
                st.rerun()
            else:
                st.error("Incorrect username or password.")
    st.stop()

_check_login()


# ── System Prompt ─────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are a commercial real estate investment analyst.
Extract all deal inputs from the provided documents and return ONLY a valid JSON object — no markdown, no explanation.

Return exactly this structure:
{
  "property": {
    "address": "string or null",
    "property_type": "string or null",
    "num_units": number or null (infer from property_type: single family/SFR=1, duplex=2, triplex=3, quadplex/fourplex=4, etc.),
    "square_feet": number or null,
    "purchase_price": number or null,
    "down_payment_pct": number (0.0-1.0, default 0.25),
    "gross_monthly_income": number or null,
    "vacancy_rate": number (0.0-1.0, default 0.0),
    "monthly_operating_expenses": number or null,
    "operating_expenses_pct": number (0.0-1.0, default 0.0),
    "monthly_property_taxes": number or null,
    "monthly_insurance": number or null,
    "monthly_utilities": number or null,
    "capex_pct": number (0.0-1.0) or null,
    "maintenance_pct": number (0.0-1.0) or null,
    "management_pct": number (0.0-1.0) or null,
    "appreciation_rate": number (0.0-1.0, default 0.03),
    "rent_growth_rate": number (0.0-1.0, default 0.02),
    "annual_cap_ex": number or null,
    "closing_costs": number or null (if not stated, default to 2% of loan_amount)
  },
  "financing": {
    "interest_rate": number (0.0-1.0, e.g. 0.07 for 7%),
    "amortization_years": number (default 30),
    "loan_amount": number or null
  },
  "missing_info": ["list of items not found in documents that affect the analysis"]
}

Rules:
- purchase_price, gross_monthly_income, and interest_rate are required for a complete analysis
- If loan_amount is not stated, calculate as purchase_price * (1 - down_payment_pct)
- If monthly_operating_expenses is stated use it; otherwise use gross_monthly_income * operating_expenses_pct
- All dollar amounts as plain numbers (no $ signs or commas)
- Interest rate and all rates as decimals (7% = 0.07)
- Return ONLY the JSON object, nothing else"""

# ── Helpers ───────────────────────────────────────────────────────────────────

def extract_json(text):
    match = re.search(r'\{[\s\S]*\}', text)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass
    return None


def get_mime(filename):
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "application/pdf"
    if lower.endswith(".xlsx") or lower.endswith(".xls"):
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if lower.endswith(".csv"):
        return "text/csv"
    if lower.endswith(".png"):
        return "image/png"
    if lower.endswith(".jpg") or lower.endswith(".jpeg"):
        return "image/jpeg"
    if lower.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if lower.endswith(".doc"):
        return "application/msword"
    return "application/octet-stream"


def extract_docx_text(file_bytes):
    """Extract plain text from a .docx file using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        # also pull table cell text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as e:
        return None


def upload_files(client, uploaded_files):
    ids = []
    for uf in uploaded_files:
        file_bytes = uf.read()
        lower = uf.name.lower()
        if lower.endswith(".docx") or lower.endswith(".doc"):
            # Convert Word doc to plain text, upload as text/plain
            text = extract_docx_text(file_bytes)
            if text:
                txt_name = uf.name + ".txt"
                buf = io.BytesIO(text.encode("utf-8"))
                resp = client.beta.files.upload(file=(txt_name, buf, "text/plain"))
                ids.append((uf.name, resp.id, "text/plain"))
            else:
                st.error("Could not extract text from " + uf.name + ". Make sure python-docx is installed: pip install python-docx")
        else:
            mime = get_mime(uf.name)
            buf = io.BytesIO(file_bytes)
            resp = client.beta.files.upload(file=(uf.name, buf, mime))
            ids.append((uf.name, resp.id, mime))
    return ids


def build_content(file_ids):
    content = []
    for fname, fid, mime in file_ids:
        if mime.startswith("image/"):
            content.append({"type": "image", "source": {"type": "file", "file_id": fid}})
        else:
            content.append({
                "type": "document",
                "source": {"type": "file", "file_id": fid},
                "title": fname,
            })
    content.append({
        "type": "text",
        "text": "Extract all deal data from these documents and return the JSON as specified.",
    })
    return content

# ── Zillow fetcher ────────────────────────────────────────────────────────────

def fetch_zillow_text(url_or_address):
    """Fetch a Zillow listing and return key property data as text for Claude."""
    try:
        import requests as _req
    except ImportError:
        return None, "The 'requests' package is required: pip install requests"

    session = _req.Session()
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
        "Sec-Fetch-Dest": "document",
        "Sec-Fetch-Mode": "navigate",
        "Sec-Fetch-Site": "none",
        "Sec-Fetch-User": "?1",
        "sec-ch-ua": '"Not A(Brand";v="99", "Google Chrome";v="121", "Chromium";v="121"',
        "sec-ch-ua-mobile": "?0",
        "sec-ch-ua-platform": '"macOS"',
        "Cache-Control": "max-age=0",
    }

    url = url_or_address.strip()
    if not url.startswith("http"):
        slug = url.replace(" ", "-").replace(",", "").replace(".", "")
        url = "https://www.zillow.com/homes/" + slug + "_rb/"

    try:
        # Seed cookies with a homepage visit first
        session.get("https://www.zillow.com/", headers=headers, timeout=10)
        headers["Referer"] = "https://www.zillow.com/"
        resp = session.get(url, headers=headers, timeout=15, allow_redirects=True)
    except Exception as e:
        return None, "Could not connect to Zillow: " + str(e)

    blocked = (
        resp.status_code in (403, 429, 503) or
        "cf-browser-verification" in resp.text or
        "Enable JavaScript and cookies" in resp.text or
        "__cf_chl" in resp.text
    )
    if blocked:
        return None, "BLOCKED"

    if resp.status_code != 200:
        return None, f"Zillow returned status {resp.status_code}."

    m = re.search(r'<script id="__NEXT_DATA__"[^>]*>(.+?)</script>', resp.text, re.DOTALL)
    if not m:
        return None, "BLOCKED"

    try:
        zdata = json.loads(m.group(1))
    except Exception:
        return None, "Could not parse Zillow page data."

    targets = {
        "price", "unformattedPrice", "bedrooms", "bathrooms", "livingArea", "lotAreaValue",
        "yearBuilt", "homeType", "propertyType", "streetAddress", "city", "state", "zipcode",
        "taxAnnualAmount", "annualHomeownersInsurance", "monthlyHoaFee", "hoaFee",
        "rentZestimate", "zestimate", "homeStatus", "description",
    }

    def find_fields(obj, depth=0):
        found = {}
        if depth > 12 or obj is None:
            return found
        if isinstance(obj, dict):
            for k, v in obj.items():
                if k in targets and v is not None and v != "" and v != 0:
                    found[k] = v
                for sk, sv in find_fields(v, depth + 1).items():
                    if sk not in found:
                        found[sk] = sv
        elif isinstance(obj, list):
            for item in obj[:3]:
                for sk, sv in find_fields(item, depth + 1).items():
                    if sk not in found:
                        found[sk] = sv
        return found

    fields = find_fields(zdata)
    if not fields:
        return None, "BLOCKED"

    label_map = {
        "price": "Listing Price", "unformattedPrice": "Price",
        "bedrooms": "Bedrooms", "bathrooms": "Bathrooms",
        "livingArea": "Square Feet", "lotAreaValue": "Lot Size",
        "yearBuilt": "Year Built", "homeType": "Property Type",
        "streetAddress": "Street Address", "city": "City",
        "state": "State", "zipcode": "Zip Code",
        "taxAnnualAmount": "Annual Property Taxes",
        "annualHomeownersInsurance": "Annual Homeowners Insurance",
        "monthlyHoaFee": "Monthly HOA Fee", "hoaFee": "HOA Fee",
        "rentZestimate": "Rent Zestimate (monthly)", "zestimate": "Zestimate",
        "homeStatus": "Status", "description": "Description",
    }
    lines = ["=== Zillow Property Listing ==="]
    for k, v in fields.items():
        lines.append(label_map.get(k, k) + ": " + str(v))

    return "\n".join(lines), None


# ── Finance math ──────────────────────────────────────────────────────────────
# calc_payment imported from dq_utils

def calc_balance(loan, annual_rate, years, elapsed_years):
    if annual_rate == 0:
        return max(0.0, loan * (1 - elapsed_years / years))
    r = annual_rate / 12
    n = years * 12
    k = int(elapsed_years * 12)
    if k >= n:
        return 0.0
    return loan * ((1 + r) ** n - (1 + r) ** k) / ((1 + r) ** n - 1)


def calc_irr(cash_flows):
    """IRR via bisection. cash_flows[0] is the initial outflow (negative)."""
    def npv(r):
        return sum(cf / (1 + r) ** t for t, cf in enumerate(cash_flows))
    try:
        lo, hi = -0.9999, 100.0
        if npv(lo) * npv(hi) > 0:
            return None
        for _ in range(300):
            mid = (lo + hi) / 2.0
            if abs(hi - lo) < 1e-10:
                break
            if npv(mid) * npv(lo) < 0:
                hi = mid
            else:
                lo = mid
        return (lo + hi) / 2.0
    except Exception:
        return None

# ── Dashboard ─────────────────────────────────────────────────────────────────
# fmt_d, fmt_p imported from dq_utils

def show_dashboard(data):
    prop    = data.get("property", {})
    fin     = data.get("financing", {})
    missing = data.get("missing_info", [])

    # ── Inputs ────────────────────────────────────────────────────────────────
    def _get(d, key, default):
        v = d.get(key)
        return v if v is not None else default

    purchase_price = _get(prop, "purchase_price", 0)
    down_pct       = _get(prop, "down_payment_pct", 0.25)
    down_payment   = purchase_price * down_pct
    loan_amount    = _get(fin, "loan_amount", None) or (purchase_price - down_payment)
    interest_rate  = _get(fin, "interest_rate", 0.07)
    amort_years    = _get(fin, "amortization_years", 30)
    closing_costs  = _get(prop, "closing_costs", None) or (loan_amount * 0.02)
    gmi            = _get(prop, "gross_monthly_income", 0)
    vacancy_rate   = _get(prop, "vacancy_rate", 0.0)
    appreciation   = _get(prop, "appreciation_rate", 0.03)
    rent_growth    = _get(prop, "rent_growth_rate", 0.02)
    annual_cap_ex  = _get(prop, "annual_cap_ex", 0)

    # Individual line items (from doc extraction or manual _ keys)
    _taxes_v = float(prop.get("monthly_property_taxes") or prop.get("_taxes") or 0)
    _ins_v   = float(prop.get("monthly_insurance")      or prop.get("_insurance") or 0)
    _util_v  = float(prop.get("monthly_utilities")      or prop.get("_utilities") or 0)
    _capex_p = float(prop.get("capex_pct")              or prop.get("_capex_pct") or 0)
    _maint_p = float(prop.get("maintenance_pct")        or prop.get("_maint_pct") or 0)
    _mgmt_p  = float(prop.get("management_pct")         or prop.get("_mgmt_pct")  or 0)
    _line_item_total = _taxes_v + _ins_v + _util_v + gmi * (_capex_p + _maint_p + _mgmt_p)

    moe = prop.get("monthly_operating_expenses")
    if moe is not None:
        monthly_opex = float(moe)
    elif _line_item_total > 0:
        monthly_opex = _line_item_total
    else:
        opex_pct     = _get(prop, "operating_expenses_pct", 0.0)
        monthly_opex = gmi * opex_pct

    # ── Core calcs ────────────────────────────────────────────────────────────
    eff_monthly     = gmi * (1 - vacancy_rate)
    monthly_cap_ex  = annual_cap_ex / 12
    monthly_noi     = eff_monthly - monthly_opex - monthly_cap_ex
    annual_noi      = monthly_noi * 12
    monthly_payment = calc_payment(loan_amount, interest_rate, amort_years)
    annual_ds       = monthly_payment * 12
    monthly_cf      = monthly_noi - monthly_payment
    annual_cf       = monthly_cf * 12
    cap_rate        = annual_noi / purchase_price if purchase_price else 0
    total_invested  = down_payment + closing_costs + annual_cap_ex
    coc_roi         = annual_cf / total_invested if total_invested else 0
    fifty_cf        = gmi * 0.5 - monthly_payment
    addr            = prop.get("address") or "Deal Analysis"
    ptype           = prop.get("property_type") or "Rental Analysis"

    # ── 30-year projections ───────────────────────────────────────────────────
    years_range = list(range(1, 31))
    pv_list, eq_list, lb_list, cf_list, mp_list, profit_list, arr_list = [], [], [], [], [], [], []
    cf_monthly_list = []
    annual_cfs = []  # store each year's cash flow for IRR
    for y in years_range:
        pv    = purchase_price * (1 + appreciation) ** y
        lb    = calc_balance(loan_amount, interest_rate, amort_years, y)
        eq    = pv - lb
        noi_y = annual_noi * (1 + rent_growth) ** (y - 1)
        cf_y  = noi_y - annual_ds
        annual_cfs.append(cf_y)
        profit = eq - pv * 0.06
        # IRR: t=0 outflow, then annual CFs, last year includes net sale proceeds
        irr_cfs = [-total_invested] + annual_cfs[:-1] + [annual_cfs[-1] + profit]
        arr = calc_irr(irr_cfs) if total_invested > 0 else None
        pv_list.append(pv)
        eq_list.append(eq)
        lb_list.append(lb)
        cf_list.append(cf_y)
        mp_list.append(monthly_payment)
        profit_list.append(profit)
        arr_list.append(arr)
        cf_monthly_list.append(cf_y / 12)

    arr_5yr = arr_list[4]

    # ── New Deal button ────────────────────────────────────────────────────────
    if st.button("＋ New Deal", type="secondary"):
        save_to_history(data)
        st.session_state.pop("dashboard_data", None)
        st.session_state["_sb_pending"] = {
            "sb_address": "", "sb_type": "", "sb_sqft": 0, "sb_price": 0,
            "sb_down": 25.0, "sb_closing": 0, "sb_gmi": 0, "sb_vacancy": 0.0,
            "sb_taxes": 0, "sb_insure": 0, "sb_util": 0,
            "sb_capex": 5.0, "sb_maint": 5.0, "sb_mgmt": 8.0,
            "sb_rate": 7.0, "sb_amort": 30, "sb_appr": 3.0, "sb_rent_g": 2.0,
            "sb_hud_token": "", "sb_bedrooms": "2 BR", "sb_units": 1,
            "zillow_input": "", "pasted_text": "",
        }
        st.rerun()

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. HERO CARD
    # ═══════════════════════════════════════════════════════════════════════════
    badge_color = "#48BB78" if monthly_cf >= 0 else "#E53E3E"
    badge_sign  = "↗" if monthly_cf >= 0 else "↘"
    st.markdown("""
    <div style="background:#F3F4F6;border:1px solid #E2E8F0;border-radius:14px;
                padding:28px 32px;margin-bottom:20px;color:#000;">
        <div style="font-size:13px;color:#666;text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;">""" + ptype + """</div>
        <div style="font-size:26px;font-weight:700;margin-bottom:12px;color:#000;">""" + addr + """</div>
        <span style="background:""" + badge_color + """;color:white;padding:6px 14px;border-radius:999px;
                     font-size:14px;font-weight:600;">
            """ + badge_sign + " " + fmt_d(monthly_cf) + """ monthly cash flow
        </span>
    </div>""", unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. CASH FLOW CHART + KEY METRICS
    # ═══════════════════════════════════════════════════════════════════════════
    col_chart, col_metrics = st.columns([3, 2])

    with col_chart:
        st.markdown('<div style="background:white;border:1px solid #E2E8F0;border-radius:10px;padding:16px;">', unsafe_allow_html=True)
        cf_fig = go.Figure()
        cf_fig.add_trace(go.Scatter(
            x=years_range, y=cf_monthly_list,
            fill="tozeroy",
            fillcolor="rgba(56,178,172,0.15)",
            line=dict(color="#0FA573", width=2.5),
            hovertemplate="Year %{x}: $%{y:,.0f}/mo<extra></extra>",
        ))
        cf_fig.update_layout(
            height=220, margin=dict(l=10, r=10, t=10, b=10),
            xaxis=dict(tickvals=[1, 2, 3, 4, 5, 10, 15, 20, 30], title="", showgrid=False, color="#888"),
            yaxis=dict(tickprefix="$", tickformat=",.0f", showgrid=True, gridcolor="#F0F0F0", color="#888"),
            plot_bgcolor="white", paper_bgcolor="white", showlegend=False,
        )
        st.plotly_chart(cf_fig, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with col_metrics:
        # eff_monthly (after vacancy) - opex - mortgage = cash flow, so the math shows correctly
        st.markdown("""
        <div style="background:white;border:1px solid #E2E8F0;border-radius:10px;padding:20px;height:100%;">
            <div style="font-size:12px;color:#888;margin-bottom:4px;">Monthly Cash Flow</div>
            <div style="font-size:32px;font-weight:700;color:#000;margin-bottom:16px;">""" + fmt_d(monthly_cf) + """ <span style="font-size:14px;color:#888;">/mo</span></div>
            <div style="display:flex;gap:20px;margin-bottom:16px;">
                <div>
                    <div style="font-size:11px;color:#888;">Eff. Income (after vacancy)</div>
                    <div style="font-size:16px;font-weight:600;color:#38A169;">""" + fmt_d(eff_monthly) + """ /mo</div>
                </div>
                <div>
                    <div style="font-size:11px;color:#888;">Total Expenses</div>
                    <div style="font-size:16px;font-weight:600;color:#ED8936;">""" + fmt_d(monthly_opex + monthly_payment) + """ /mo</div>
                </div>
            </div>
            <div style="display:flex;gap:20px;">
                <div>
                    <div style="font-size:11px;color:#888;">5-yr Ann. Return</div>
                    <div style="font-size:16px;font-weight:700;color:#000;">""" + (fmt_p(arr_5yr) if arr_5yr else "—") + """</div>
                </div>
                <div>
                    <div style="font-size:11px;color:#888;">Mortgage Pmt</div>
                    <div style="font-size:16px;font-weight:700;color:#000;">""" + fmt_d(monthly_payment) + """</div>
                </div>
            </div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. RETURNS + 50% RULE METRICS
    # ═══════════════════════════════════════════════════════════════════════════
    st.markdown('<div style="background:white;border:1px solid #E2E8F0;border-radius:10px;padding:20px;margin-bottom:16px;">', unsafe_allow_html=True)
    st.markdown('<div style="font-weight:700;font-size:16px;margin-bottom:14px;color:#000;">Returns</div>', unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("NOI", fmt_d(annual_noi))
    c2.metric("CoC ROI", fmt_p(coc_roi))
    c3.metric("Pro Forma Cap", fmt_p(cap_rate))
    c4.metric("Purchase Cap", fmt_p(cap_rate))
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. INCOME & EXPENSE BREAKDOWN
    # ═══════════════════════════════════════════════════════════════════════════
    st.markdown('<div style="background:white;border:1px solid #E2E8F0;border-radius:10px;padding:20px;margin-bottom:16px;">', unsafe_allow_html=True)
    st.markdown('<div style="font-weight:700;font-size:15px;color:#000;margin-bottom:14px;">Income & Expense Breakdown</div>', unsafe_allow_html=True)

    col_inc, col_exp, col_cf = st.columns(3)

    def breakdown_row(label, value, indent=False, bold=False, color="#000", suffix=""):
        pad = "18px" if indent else "0px"
        weight = "700" if bold else "400"
        return (
            '<div style="display:flex;justify-content:space-between;padding:5px 0;'
            'padding-left:' + pad + ';border-bottom:1px solid #F7F7F7;">'
            '<span style="font-size:13px;color:' + color + ';font-weight:' + weight + ';">' + label + '</span>'
            '<span style="font-size:13px;font-weight:' + weight + ';color:' + color + ';">' + suffix + fmt_d(value) + '</span>'
            '</div>'
        )

    vacancy_mo = gmi * vacancy_rate

    with col_inc:
        st.markdown('<div style="font-weight:600;font-size:12px;color:#777;margin-bottom:8px;text-transform:uppercase;letter-spacing:0.5px;">Income</div>', unsafe_allow_html=True)
        st.markdown(breakdown_row("Gross Monthly Income", gmi), unsafe_allow_html=True)
        st.markdown(breakdown_row("Vacancy (" + str(round(vacancy_rate * 100, 1)) + "%)", vacancy_mo, indent=True, color="#E53E3E", suffix="− "), unsafe_allow_html=True)
        st.markdown(breakdown_row("= Effective Income", eff_monthly, bold=True, color="#38A169"), unsafe_allow_html=True)

    with col_exp:
        st.markdown('<div style="font-weight:600;font-size:12px;color:#777;margin-bottom:8px;text-transform:uppercase;letter-spacing:0.5px;">Operating Expenses</div>', unsafe_allow_html=True)
        d_prop = st.session_state.get("dashboard_data", {}).get("property", {})
        taxes_v = float(d_prop.get("monthly_property_taxes") or d_prop.get("_taxes") or 0)
        ins_v   = float(d_prop.get("monthly_insurance")      or d_prop.get("_insurance") or 0)
        util_v  = float(d_prop.get("monthly_utilities")      or d_prop.get("_utilities") or 0)
        capex_p = float(d_prop.get("capex_pct")              or d_prop.get("_capex_pct") or 0)
        maint_p = float(d_prop.get("maintenance_pct")        or d_prop.get("_maint_pct") or 0)
        mgmt_p  = float(d_prop.get("management_pct")         or d_prop.get("_mgmt_pct")  or 0)
        has_detail = any([taxes_v, ins_v, util_v, capex_p, maint_p, mgmt_p])

        if has_detail:
            if taxes_v:
                st.markdown(breakdown_row("Property Taxes", taxes_v, indent=True), unsafe_allow_html=True)
            if ins_v:
                st.markdown(breakdown_row("Insurance", ins_v, indent=True), unsafe_allow_html=True)
            if util_v:
                st.markdown(breakdown_row("Utilities", util_v, indent=True), unsafe_allow_html=True)
            if capex_p:
                st.markdown(breakdown_row("CapEx (" + str(round(capex_p * 100, 1)) + "% rent)", gmi * capex_p, indent=True), unsafe_allow_html=True)
            if maint_p:
                st.markdown(breakdown_row("Maintenance (" + str(round(maint_p * 100, 1)) + "% rent)", gmi * maint_p, indent=True), unsafe_allow_html=True)
            if mgmt_p:
                st.markdown(breakdown_row("Management (" + str(round(mgmt_p * 100, 1)) + "% rent)", gmi * mgmt_p, indent=True), unsafe_allow_html=True)
        else:
            st.markdown(breakdown_row("Operating Expenses", monthly_opex, indent=True), unsafe_allow_html=True)

        if monthly_cap_ex > 0:
            st.markdown(breakdown_row("CapEx Reserve", monthly_cap_ex, indent=True), unsafe_allow_html=True)
        st.markdown(breakdown_row("= Total Op. Expenses", monthly_opex + monthly_cap_ex, bold=True, color="#ED8936"), unsafe_allow_html=True)

    with col_cf:
        st.markdown('<div style="font-weight:600;font-size:12px;color:#777;margin-bottom:8px;text-transform:uppercase;letter-spacing:0.5px;">Cash Flow</div>', unsafe_allow_html=True)
        st.markdown(breakdown_row("Effective Income", eff_monthly, color="#38A169"), unsafe_allow_html=True)
        st.markdown(breakdown_row("Op. Expenses", monthly_opex + monthly_cap_ex, indent=True, color="#E53E3E", suffix="− "), unsafe_allow_html=True)
        st.markdown(breakdown_row("= NOI", monthly_noi, bold=True), unsafe_allow_html=True)
        st.markdown(breakdown_row("Mortgage (P&I)", monthly_payment, indent=True, color="#E53E3E", suffix="− "), unsafe_allow_html=True)
        cf_color = "#38A169" if monthly_cf >= 0 else "#E53E3E"
        st.markdown(breakdown_row("= Monthly Cash Flow", monthly_cf, bold=True, color=cf_color), unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. EXPENSE BREAKDOWN (donut) + DEAL SUMMARY
    # ═══════════════════════════════════════════════════════════════════════════
    col_donut, col_summary = st.columns([2, 3])

    with col_donut:
        st.markdown('<div style="background:white;border:1px solid #E2E8F0;border-radius:10px;padding:20px;">', unsafe_allow_html=True)
        st.markdown('<div style="font-weight:700;font-size:15px;color:#000;margin-bottom:12px;">Monthly Expense Breakdown</div>', unsafe_allow_html=True)
        vacancy_mo   = gmi * vacancy_rate
        labels  = ["Mortgage", "Operating Exp", "Vacancy", "CapEx"]
        values  = [monthly_payment, monthly_opex, vacancy_mo, monthly_cap_ex]
        colors  = ["#4299E1", "#ED8936", "#ECC94B", "#9F7AEA"]
        # filter out zero slices
        lv = [(l, v, c) for l, v, c in zip(labels, values, colors) if v > 0]
        if lv:
            dl, dv, dc = zip(*lv)
            donut = go.Figure(go.Pie(
                labels=list(dl), values=list(dv),
                hole=0.55,
                marker=dict(colors=list(dc)),
                textinfo="none",
                hovertemplate="%{label}: $%{value:,.0f}<extra></extra>",
            ))
            total_exp = sum(dv)
            donut.add_annotation(text=fmt_d(total_exp), x=0.5, y=0.5,
                                 font=dict(size=16, color="#000", family="sans-serif"),
                                 showarrow=False)
            donut.update_layout(
                height=220, margin=dict(l=0, r=0, t=0, b=0),
                plot_bgcolor="white", paper_bgcolor="white",
                legend=dict(orientation="h", x=0, y=-0.1, font=dict(size=11)),
                showlegend=True,
            )
            st.plotly_chart(donut, use_container_width=True)
        for l, v, c in zip(labels, values, colors):
            if v <= 0:
                continue
            st.markdown(
                '<div style="display:flex;justify-content:space-between;font-size:13px;padding:3px 0;color:#000;">'
                '<span><span style="color:' + c + ';margin-right:6px;">●</span>' + l + '</span>'
                '<span style="font-weight:600;">' + fmt_d(v) + '</span></div>',
                unsafe_allow_html=True,
            )
        st.markdown("</div>", unsafe_allow_html=True)

    with col_summary:
        st.markdown('<div style="background:white;border:1px solid #E2E8F0;border-radius:10px;padding:20px;">', unsafe_allow_html=True)
        st.markdown('<div style="font-weight:700;font-size:15px;color:#000;margin-bottom:14px;">Deal Summary</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        c1.metric("Purchase Price", fmt_d(purchase_price))
        c2.metric("Down Payment", fmt_d(down_payment) + " (" + str(round(down_pct * 100)) + "%)")
        c1, c2 = st.columns(2)
        c1.metric("Loan Amount", fmt_d(loan_amount))
        c2.metric("Interest Rate", fmt_p(interest_rate))
        st.markdown('<div style="font-weight:700;font-size:15px;color:#000;margin:16px 0 14px 0;">50% Rule</div>', unsafe_allow_html=True)
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Total Mo. Income", fmt_d(gmi))
        c2.metric("50% for Expenses", fmt_d(gmi * 0.5))
        c3.metric("Monthly P&I", fmt_d(monthly_payment))
        c4.metric("50% Rule Cash Flow", fmt_d(fifty_cf))
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. 30-YEAR PROJECTION CHART + TABLE
    # ═══════════════════════════════════════════════════════════════════════════
    st.markdown('<div style="background:white;border:1px solid #E2E8F0;border-radius:10px;padding:20px;">', unsafe_allow_html=True)
    st.markdown('<div style="font-weight:700;font-size:15px;color:#000;margin-bottom:12px;">30-Year Projections</div>', unsafe_allow_html=True)

    proj_fig = go.Figure()
    proj_fig.add_trace(go.Scatter(
        x=years_range, y=pv_list, name="Property Value",
        line=dict(color="#2F5496", width=2.5),
        hovertemplate="Year %{x}: $%{y:,.0f}<extra>Property Value</extra>",
    ))
    proj_fig.add_trace(go.Scatter(
        x=years_range, y=eq_list, name="Equity",
        line=dict(color="#0FA573", width=2.5),
        hovertemplate="Year %{x}: $%{y:,.0f}<extra>Equity</extra>",
    ))
    proj_fig.add_trace(go.Scatter(
        x=years_range, y=lb_list, name="Loan Balance",
        line=dict(color="#ED8936", width=2, dash="dash"),
        hovertemplate="Year %{x}: $%{y:,.0f}<extra>Loan Balance</extra>",
    ))
    proj_fig.update_layout(
        height=300, margin=dict(l=10, r=10, t=10, b=10),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        yaxis=dict(tickprefix="$", tickformat=",.0f", gridcolor="#F0F0F0"),
        xaxis=dict(title="Year", gridcolor="#F0F0F0"),
        plot_bgcolor="white", paper_bgcolor="white", hovermode="x unified",
    )
    st.plotly_chart(proj_fig, use_container_width=True)

    show_yrs = [1, 2, 3, 4, 5, 10, 15, 20, 30]
    idx = {y: i for i, y in enumerate(years_range)}
    rows = {
        "Property Value": [], "Equity": [], "Loan Balance": [],
        "Cash Flow":      [], "Mortgage Pmt": [], "Profit if Sold": [], "Ann. Return": [],
    }
    for sy in show_yrs:
        i = idx[sy]
        rows["Property Value"].append(fmt_d(pv_list[i], short=True))
        rows["Equity"].append(fmt_d(eq_list[i], short=True))
        rows["Loan Balance"].append(fmt_d(lb_list[i], short=True))
        rows["Cash Flow"].append(fmt_d(cf_list[i], short=True))
        rows["Mortgage Pmt"].append(fmt_d(mp_list[i]))
        rows["Profit if Sold"].append(fmt_d(profit_list[i], short=True))
        a = arr_list[i]
        rows["Ann. Return"].append(fmt_p(a) if a is not None else "—")

    row_labels = list(rows.keys())
    col_labels = ["Year " + str(y) for y in show_yrs]

    header_cells = "".join(
        '<th style="padding:8px 12px;text-align:right;font-size:12px;font-weight:600;'
        'color:#555;border-bottom:2px solid #E2E8F0;white-space:nowrap;">' + c + '</th>'
        for c in col_labels
    )
    html_rows = ""
    for ri, label in enumerate(row_labels):
        bg = "#ffffff" if ri % 2 == 0 else "#F9FAFB"
        cells = "".join(
            '<td style="padding:8px 12px;text-align:right;font-size:13px;'
            'color:#000;white-space:nowrap;">' + rows[label][ci] + '</td>'
            for ci in range(len(show_yrs))
        )
        html_rows += (
            '<tr style="background:' + bg + ';">'
            '<td style="padding:8px 12px;font-size:13px;font-weight:600;color:#000;'
            'white-space:nowrap;border-right:1px solid #E2E8F0;">' + label + '</td>'
            + cells + '</tr>'
        )

    table_html = (
        '<div style="overflow-x:auto;margin-top:12px;">'
        '<table style="width:100%;border-collapse:collapse;font-family:sans-serif;">'
        '<thead><tr style="background:#F3F4F6;">'
        '<th style="padding:8px 12px;text-align:left;font-size:12px;font-weight:600;'
        'color:#555;border-bottom:2px solid #E2E8F0;border-right:1px solid #E2E8F0;"></th>'
        + header_cells +
        '</tr></thead>'
        '<tbody>' + html_rows + '</tbody>'
        '</table></div>'
    )
    st.markdown(table_html, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Missing info ──────────────────────────────────────────────────────────
    if missing:
        items = "<br>".join("• " + m for m in missing)
        st.markdown(
            '<div class="missing-box"><b>⚠️ Missing info / assumptions used:</b><br>' + items + '</div>',
            unsafe_allow_html=True,
        )


# ── Sidebar state helpers ─────────────────────────────────────────────────────

# Apply any staged sidebar values from document extraction (must run before widgets render)
if "_sb_pending" in st.session_state:
    for _k, _v in st.session_state.pop("_sb_pending").items():
        st.session_state[_k] = _v

_SB_DEFAULTS = {
    "sb_address": "", "sb_type": "", "sb_sqft": 0, "sb_price": 0,
    "sb_down": 25.0, "sb_closing": 0, "sb_gmi": 0, "sb_vacancy": 0.0,
    "sb_taxes": 0, "sb_insure": 0, "sb_util": 0,
    "sb_capex": 5.0, "sb_maint": 5.0, "sb_mgmt": 8.0,
    "sb_rate": 7.0, "sb_amort": 30, "sb_appr": 3.0, "sb_rent_g": 2.0,
    "sb_hud_token": "", "sb_bedrooms": "2 BR", "sb_units": 1,
}
for _k, _v in _SB_DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ── Tax & insurance estimator ─────────────────────────────────────────────────

# County-level effective property tax rates (GEOID = 5-digit state+county FIPS)
_COUNTY_TAX_RATE = {
    # Texas
    "48201":0.0170,"48113":0.0195,"48439":0.0185,"48029":0.0180,"48453":0.0175,
    "48085":0.0185,"48121":0.0185,"48141":0.0160,"48215":0.0155,"48157":0.0175,
    "48245":0.0185,"48491":0.0175,"48167":0.0175,"48355":0.0170,"48339":0.0160,
    # California
    "06037":0.0110,"06073":0.0080,"06059":0.0080,"06065":0.0105,"06071":0.0100,
    "06001":0.0080,"06067":0.0080,"06013":0.0080,"06019":0.0075,"06075":0.0055,
    "06085":0.0078,"06077":0.0080,"06029":0.0070,"06083":0.0080,"06111":0.0075,
    # Florida
    "12086":0.0100,"12011":0.0100,"12099":0.0105,"12057":0.0100,"12095":0.0090,
    "12103":0.0090,"12031":0.0100,"12071":0.0100,"12105":0.0095,"12009":0.0090,
    "12097":0.0095,"12021":0.0095,"12117":0.0090,"12127":0.0090,"12053":0.0090,
    # New York
    "36047":0.0075,"36081":0.0090,"36005":0.0130,"36061":0.0085,"36103":0.0180,
    "36059":0.0220,"36119":0.0200,"36029":0.0250,"36055":0.0240,"36067":0.0240,
    "36001":0.0220,"36007":0.0250,"36013":0.0230,"36063":0.0250,"36065":0.0220,
    # Illinois
    "17031":0.0220,"17043":0.0195,"17097":0.0230,"17197":0.0230,"17089":0.0235,
    "17111":0.0240,"17037":0.0220,"17019":0.0215,"17093":0.0220,"17115":0.0230,
    # Pennsylvania
    "42101":0.0310,"42003":0.0160,"42091":0.0170,"42017":0.0150,"42045":0.0220,
    "42029":0.0150,"42049":0.0220,"42065":0.0170,"42071":0.0160,"42077":0.0160,
    # Georgia
    "13121":0.0100,"13089":0.0100,"13067":0.0075,"13135":0.0090,"13063":0.0110,
    "13057":0.0095,"13117":0.0095,"13151":0.0090,"13247":0.0095,"13045":0.0085,
    # North Carolina
    "37119":0.0085,"37183":0.0078,"37081":0.0090,"37067":0.0090,"37063":0.0100,
    "37051":0.0080,"37071":0.0085,"37097":0.0080,"37179":0.0080,"37035":0.0080,
    # Michigan
    "26163":0.0220,"26125":0.0145,"26099":0.0135,"26081":0.0145,"26161":0.0160,
    "26049":0.0150,"26065":0.0140,"26093":0.0150,"26021":0.0140,"26145":0.0145,
    # Ohio
    "39035":0.0200,"39049":0.0160,"39061":0.0155,"39153":0.0200,"39113":0.0190,
    "39095":0.0200,"39029":0.0180,"39057":0.0175,"39099":0.0195,"39055":0.0185,
    # Arizona
    "04013":0.0060,"04019":0.0065,"04021":0.0060,"04025":0.0060,"04007":0.0062,
    "04015":0.0060,"04027":0.0065,"04005":0.0060,"04017":0.0060,"04009":0.0060,
    # Colorado
    "08031":0.0055,"08059":0.0050,"08005":0.0055,"08001":0.0050,"08041":0.0050,
    "08069":0.0055,"08013":0.0055,"08123":0.0050,"08035":0.0055,"08014":0.0050,
    # Washington
    "53033":0.0095,"53053":0.0100,"53061":0.0095,"53063":0.0095,"53011":0.0090,
    "53035":0.0095,"53057":0.0090,"53067":0.0090,"53077":0.0090,"53047":0.0095,
    # Tennessee
    "47157":0.0135,"47037":0.0075,"47093":0.0055,"47065":0.0065,"47187":0.0070,
    "47009":0.0060,"47149":0.0065,"47165":0.0075,"47001":0.0055,"47037":0.0075,
    # Nevada
    "32003":0.0055,"32031":0.0055,"32029":0.0055,"32510":0.0055,
    # Oregon
    "41051":0.0090,"41067":0.0085,"41039":0.0085,"41005":0.0085,"41047":0.0085,
    "41043":0.0085,"41017":0.0080,"41029":0.0085,"41071":0.0080,"41053":0.0085,
    # Massachusetts
    "25025":0.0075,"25017":0.0105,"25021":0.0110,"25027":0.0135,"25009":0.0130,
    "25013":0.0190,"25015":0.0120,"25023":0.0130,"25005":0.0100,"25011":0.0130,
    # Minnesota
    "27053":0.0110,"27123":0.0125,"27037":0.0110,"27003":0.0125,"27163":0.0110,
    "27019":0.0110,"27139":0.0110,"27049":0.0110,"27109":0.0120,"27145":0.0110,
    # Wisconsin
    "55079":0.0220,"55025":0.0165,"55133":0.0185,"55009":0.0180,"55139":0.0175,
    "55087":0.0180,"55105":0.0175,"55059":0.0175,"55021":0.0180,"55071":0.0175,
    # Missouri
    "29189":0.0110,"29095":0.0100,"29510":0.0145,"29099":0.0095,"29183":0.0105,
    "29019":0.0100,"29047":0.0100,"29031":0.0100,"29077":0.0095,"29113":0.0095,
    # Indiana
    "18097":0.0100,"18057":0.0105,"18089":0.0120,"18003":0.0090,"18141":0.0090,
    "18163":0.0100,"18095":0.0100,"18035":0.0095,"18019":0.0095,"18011":0.0090,
    # Louisiana
    "22071":0.0060,"22051":0.0050,"22033":0.0055,"22103":0.0055,"22017":0.0050,
    "22019":0.0055,"22015":0.0050,"22055":0.0050,"22047":0.0050,"22109":0.0050,
    # South Carolina
    "45019":0.0055,"45079":0.0055,"45051":0.0040,"45045":0.0055,"45007":0.0050,
    "45057":0.0055,"45063":0.0055,"45091":0.0050,"45083":0.0050,"45037":0.0050,
    # Virginia
    "51059":0.0105,"51153":0.0110,"51810":0.0100,"51107":0.0090,"51013":0.0095,
    "51041":0.0085,"51087":0.0080,"51760":0.0095,"51199":0.0085,"51085":0.0080,
    # Maryland
    "24033":0.0130,"24031":0.0090,"24003":0.0090,"24005":0.0105,"24510":0.0230,
    "24027":0.0090,"24021":0.0095,"24043":0.0095,"24015":0.0085,"24013":0.0085,
    # New Jersey
    "34003":0.0205,"34013":0.0280,"34017":0.0200,"34023":0.0210,"34025":0.0215,
    "34027":0.0200,"34039":0.0240,"34007":0.0230,"34031":0.0240,"34019":0.0210,
    "34035":0.0200,"34029":0.0230,"34041":0.0210,"34037":0.0200,"34021":0.0220,
    # Connecticut
    "09003":0.0200,"09009":0.0200,"09001":0.0150,"09007":0.0180,"09011":0.0190,
    "09005":0.0195,"09013":0.0195,"09015":0.0195,
    # Utah
    "49035":0.0055,"49049":0.0055,"49011":0.0055,"49057":0.0060,"49053":0.0055,
    "49043":0.0055,"49003":0.0055,"49005":0.0055,"49023":0.0055,
    # Idaho
    "16001":0.0060,"16027":0.0065,"16055":0.0060,"16069":0.0060,"16005":0.0060,
    # New Mexico
    "35001":0.0070,"35013":0.0075,"35049":0.0055,"35043":0.0060,"35045":0.0060,
    # Hawaii
    "15003":0.0030,"15009":0.0030,"15007":0.0030,"15005":0.0030,
    # Oklahoma
    "40109":0.0110,"40143":0.0100,"40027":0.0100,"40017":0.0100,"40083":0.0100,
    # Kentucky
    "21111":0.0080,"21067":0.0075,"21059":0.0080,"21019":0.0080,"21029":0.0075,
    # Alabama
    "01073":0.0050,"01089":0.0045,"01097":0.0045,"01101":0.0040,"01077":0.0040,
    # Iowa
    "19153":0.0155,"19113":0.0160,"19013":0.0155,"19163":0.0150,"19049":0.0150,
    # Kansas
    "20091":0.0125,"20173":0.0125,"20209":0.0135,"20045":0.0120,"20015":0.0120,
    # Nebraska
    "31055":0.0175,"31109":0.0185,"31153":0.0175,"31019":0.0175,"31079":0.0175,
    # Alaska
    "02020":0.0120,"02090":0.0110,"02170":0.0100,
    # Rhode Island
    "44007":0.0160,"44003":0.0150,"44005":0.0150,"44009":0.0155,
    # Delaware
    "10003":0.0060,"10001":0.0055,"10005":0.0055,
    # New Hampshire
    "33011":0.0200,"33015":0.0200,"33013":0.0195,"33005":0.0195,
    # Maine
    "23005":0.0115,"23031":0.0105,"23001":0.0110,"23019":0.0110,
    # Montana
    "30111":0.0080,"30063":0.0080,"30013":0.0080,"30049":0.0075,
    # South Dakota
    "46099":0.0110,"46083":0.0110,"46029":0.0110,
    # North Dakota
    "38017":0.0090,"38015":0.0085,"38059":0.0085,
    # West Virginia
    "54039":0.0055,"54011":0.0055,"54107":0.0055,
    # Mississippi
    "28049":0.0080,"28047":0.0065,"28059":0.0070,
    # Arkansas
    "05119":0.0065,"05007":0.0055,"05143":0.0055,"05131":0.0055,
    # Wyoming
    "56025":0.0055,"56021":0.0060,"56013":0.0055,
}

# County-level annual insurance estimates (GEOID → annual premium $)
_COUNTY_INSURANCE = {
    # High-risk Florida coastal
    "12086":3200,"12011":2800,"12099":3000,"12057":2600,"12095":2400,
    "12103":2500,"12031":2400,"12071":3000,"12009":2400,"12097":2800,
    # Texas (hurricane/hail zone)
    "48201":2800,"48113":2900,"48439":2700,"48029":2600,"48453":2500,
    "48085":2700,"48121":2700,"48215":2600,"48157":2600,
    # Oklahoma/Kansas (tornado alley)
    "40109":2600,"40143":2500,"20173":2400,"20091":2300,
    # Louisiana (hurricane)
    "22071":3200,"22051":3000,"22033":3000,"22103":2800,
    # High-cost metros (CA)
    "06037":1600,"06075":1800,"06073":1400,"06059":1400,"06085":1600,
    # Mid-Atlantic
    "24510":1300,"42101":1200,"34013":1400,"34039":1500,
    # Northeast
    "36047":1800,"36081":1700,"36061":2000,"36059":1600,"36119":1600,
    "25025":1700,"25017":1600,
    # Pacific Northwest (low risk)
    "53033":1100,"41051":900,"41067":900,
    # Mountain West (low risk)
    "08031":1900,"04013":1300,"49035":900,
}


def fetch_census_county(address):
    """Use Census Geocoder to get county name and 5-digit GEOID from address.
    Returns (geoid, county_name, state_fips) or (None, None, None)."""
    try:
        import requests as _req
    except ImportError:
        return None, None, None
    try:
        resp = _req.get(
            "https://geocoding.geo.census.gov/geocoder/geographies/onelineaddress",
            params={
                "address": address,
                "benchmark": "Public_AR_Current",
                "vintage": "Current_Current",
                "layers": "Counties",
                "format": "json",
            },
            timeout=10,
        )
        if resp.status_code != 200:
            return None, None, None
        matches = resp.json().get("result", {}).get("addressMatches", [])
        if not matches:
            return None, None, None
        counties = matches[0].get("geographies", {}).get("Counties", [])
        if not counties:
            return None, None, None
        c = counties[0]
        geoid      = c.get("GEOID")       # 5-digit e.g. "48201"
        county_name = c.get("NAME")        # e.g. "Harris County"
        state_fips  = c.get("STATE")       # e.g. "48"
        return geoid, county_name, state_fips
    except Exception:
        return None, None, None


# Effective property tax rates by state (% of assessed value, 2024 averages)
_STATE_TAX_RATE = {
    "AL":0.0039,"AK":0.0104,"AZ":0.0051,"AR":0.0057,"CA":0.0071,"CO":0.0049,
    "CT":0.0167,"DE":0.0055,"FL":0.0083,"GA":0.0081,"HI":0.0026,"ID":0.0063,
    "IL":0.0188,"IN":0.0081,"IA":0.0150,"KS":0.0124,"KY":0.0083,"LA":0.0055,
    "ME":0.0109,"MD":0.0100,"MA":0.0105,"MI":0.0138,"MN":0.0102,"MS":0.0066,
    "MO":0.0091,"MT":0.0074,"NE":0.0154,"NV":0.0055,"NH":0.0186,"NJ":0.0223,
    "NM":0.0067,"NY":0.0140,"NC":0.0070,"ND":0.0088,"OH":0.0137,"OK":0.0085,
    "OR":0.0082,"PA":0.0136,"RI":0.0130,"SC":0.0052,"SD":0.0108,"TN":0.0056,
    "TX":0.0160,"UT":0.0052,"VT":0.0178,"VA":0.0075,"WA":0.0084,"WV":0.0055,
    "WI":0.0151,"WY":0.0055,"DC":0.0055,
}

# Average annual homeowners insurance premium by state (2024 estimates)
_STATE_INSURANCE = {
    "AL":1600,"AK":990,"AZ":1200,"AR":1700,"CA":1200,"CO":1700,"CT":1400,
    "DE":900,"FL":2100,"GA":1500,"HI":520,"ID":900,"IL":1400,"IN":1200,
    "IA":1100,"KS":2100,"KY":1600,"LA":2400,"ME":900,"MD":1100,"MA":1400,
    "MI":1100,"MN":1400,"MS":1900,"MO":1700,"MT":1400,"NE":1800,"NV":900,
    "NH":1000,"NJ":1200,"NM":1400,"NY":1300,"NC":1500,"ND":1500,"OH":1100,
    "OK":2300,"OR":800,"PA":1000,"RI":1400,"SC":1600,"SD":1500,"TN":1600,
    "TX":2300,"UT":800,"VT":900,"VA":1100,"WA":900,"WV":1000,"WI":900,
    "WY":1300,"DC":1100,
}

_STATE_NAMES = {
    "alabama":"AL","alaska":"AK","arizona":"AZ","arkansas":"AR","california":"CA",
    "colorado":"CO","connecticut":"CT","delaware":"DE","florida":"FL","georgia":"GA",
    "hawaii":"HI","idaho":"ID","illinois":"IL","indiana":"IN","iowa":"IA","kansas":"KS",
    "kentucky":"KY","louisiana":"LA","maine":"ME","maryland":"MD","massachusetts":"MA",
    "michigan":"MI","minnesota":"MN","mississippi":"MS","missouri":"MO","montana":"MT",
    "nebraska":"NE","nevada":"NV","new hampshire":"NH","new jersey":"NJ","new mexico":"NM",
    "new york":"NY","north carolina":"NC","north dakota":"ND","ohio":"OH","oklahoma":"OK",
    "oregon":"OR","pennsylvania":"PA","rhode island":"RI","south carolina":"SC",
    "south dakota":"SD","tennessee":"TN","texas":"TX","utah":"UT","vermont":"VT",
    "virginia":"VA","washington":"WA","west virginia":"WV","wisconsin":"WI","wyoming":"WY",
    "washington dc":"DC","district of columbia":"DC",
}


# Average monthly utilities per unit by state (2024 estimates)
_STATE_UTILITIES = {
    "AL":180,"AK":350,"AZ":160,"AR":170,"CA":130,"CO":140,"CT":210,"DE":170,
    "FL":200,"GA":185,"HI":280,"ID":130,"IL":175,"IN":175,"IA":160,"KS":165,
    "KY":175,"LA":200,"ME":210,"MD":175,"MA":200,"MI":185,"MN":185,"MS":200,
    "MO":175,"MT":150,"NE":165,"NV":150,"NH":200,"NJ":185,"NM":145,"NY":190,
    "NC":175,"ND":170,"OH":175,"OK":175,"OR":120,"PA":180,"RI":195,"SC":195,
    "SD":165,"TN":185,"TX":190,"UT":130,"VT":200,"VA":170,"WA":120,"WV":175,
    "WI":170,"WY":155,"DC":165,
}

# Gross rental yield by state (annual rent / purchase price, 2024 estimates)
_STATE_RENT_YIELD = {
    "AL":0.075,"AK":0.055,"AZ":0.058,"AR":0.078,"CA":0.042,"CO":0.052,"CT":0.060,
    "DE":0.062,"FL":0.065,"GA":0.070,"HI":0.045,"ID":0.058,"IL":0.072,"IN":0.085,
    "IA":0.080,"KS":0.082,"KY":0.078,"LA":0.072,"ME":0.058,"MD":0.060,"MA":0.048,
    "MI":0.085,"MN":0.065,"MS":0.085,"MO":0.078,"MT":0.055,"NE":0.075,"NV":0.058,
    "NH":0.055,"NJ":0.052,"NM":0.065,"NY":0.045,"NC":0.065,"ND":0.075,"OH":0.090,
    "OK":0.085,"OR":0.052,"PA":0.070,"RI":0.058,"SC":0.068,"SD":0.075,"TN":0.072,
    "TX":0.068,"UT":0.055,"VT":0.055,"VA":0.058,"WA":0.050,"WV":0.085,"WI":0.072,
    "WY":0.060,"DC":0.042,
}

# Rent per sq ft per month by state (2024 estimates)
_STATE_RENT_SQFT = {
    "AL":0.85,"AK":1.20,"AZ":1.15,"AR":0.80,"CA":2.20,"CO":1.40,"CT":1.50,
    "DE":1.30,"FL":1.45,"GA":1.20,"HI":2.50,"ID":1.10,"IL":1.30,"IN":0.90,
    "IA":0.85,"KS":0.90,"KY":0.85,"LA":1.00,"ME":1.10,"MD":1.55,"MA":2.00,
    "MI":0.95,"MN":1.10,"MS":0.85,"MO":0.90,"MT":1.10,"NE":0.95,"NV":1.25,
    "NH":1.30,"NJ":1.80,"NM":1.05,"NY":2.50,"NC":1.15,"ND":0.95,"OH":0.85,
    "OK":0.85,"OR":1.40,"PA":1.10,"RI":1.45,"SC":1.10,"SD":0.90,"TN":1.15,
    "TX":1.20,"UT":1.20,"VT":1.20,"VA":1.40,"WA":1.60,"WV":0.75,"WI":0.95,
    "WY":0.90,"DC":2.80,
}

# City-level data: city_lower → (state, tax_rate, median_monthly_rent, annual_insurance, rent_per_sqft)
_CITY_DATA = {
    "new york city":("NY",0.0140,3400,1800,2.80),"new york":("NY",0.0140,3400,1800,2.80),
    "nyc":("NY",0.0140,3400,1800,2.80),"manhattan":("NY",0.0200,4200,2000,4.00),
    "brooklyn":("NY",0.0140,3000,1700,2.50),"bronx":("NY",0.0140,2200,1600,1.80),
    "queens":("NY",0.0140,2600,1700,2.00),"boston":("MA",0.0105,2800,1600,2.40),
    "cambridge":("MA",0.0105,3200,1600,2.80),"philadelphia":("PA",0.0310,1600,1100,1.20),
    "pittsburgh":("PA",0.0150,1100,1000,0.90),"baltimore":("MD",0.1100,1400,1100,1.10),
    "washington dc":("DC",0.0085,2400,1200,2.00),"dc":("DC",0.0085,2400,1200,2.00),
    "newark":("NJ",0.0260,1600,1300,1.30),"jersey city":("NJ",0.0200,2400,1300,2.00),
    "hartford":("CT",0.0320,1300,1400,1.00),"providence":("RI",0.0180,1400,1400,1.10),
    "miami":("FL",0.0100,2400,2500,1.90),"miami beach":("FL",0.0100,2800,2700,2.20),
    "orlando":("FL",0.0085,1700,2200,1.40),"tampa":("FL",0.0095,1800,2300,1.50),
    "jacksonville":("FL",0.0095,1500,2100,1.20),"fort lauderdale":("FL",0.0100,2000,2500,1.60),
    "st. petersburg":("FL",0.0090,1700,2200,1.40),"cape coral":("FL",0.0090,1600,2400,1.30),
    "atlanta":("GA",0.0110,1700,1600,1.40),"charlotte":("NC",0.0082,1600,1700,1.30),
    "raleigh":("NC",0.0078,1500,1600,1.20),"durham":("NC",0.0090,1400,1600,1.10),
    "nashville":("TN",0.0072,1800,1800,1.50),"memphis":("TN",0.0140,1100,2000,0.85),
    "richmond":("VA",0.0120,1500,1100,1.20),"virginia beach":("VA",0.0099,1400,1200,1.10),
    "new orleans":("LA",0.0060,1400,2700,1.10),"baton rouge":("LA",0.0055,1200,2500,0.95),
    "birmingham":("AL",0.0055,1100,1700,0.85),"charleston":("SC",0.0055,1800,1900,1.50),
    "chicago":("IL",0.0220,1800,1600,1.50),"indianapolis":("IN",0.0095,1100,1200,0.90),
    "columbus":("OH",0.0155,1200,1100,0.95),"cleveland":("OH",0.0170,900,1100,0.75),
    "cincinnati":("OH",0.0145,1100,1100,0.90),"detroit":("MI",0.0200,1000,1100,0.80),
    "grand rapids":("MI",0.0145,1200,1100,0.95),"minneapolis":("MN",0.0110,1500,1400,1.20),
    "st. paul":("MN",0.0120,1400,1400,1.10),"milwaukee":("WI",0.0210,1200,1000,0.95),
    "madison":("WI",0.0160,1400,1000,1.10),"st. louis":("MO",0.0120,1100,1800,0.90),
    "kansas city":("MO",0.0110,1200,1800,0.95),"omaha":("NE",0.0180,1100,1900,0.90),
    "des moines":("IA",0.0160,1100,1200,0.85),"dallas":("TX",0.0200,1600,2500,1.30),
    "fort worth":("TX",0.0200,1400,2500,1.15),"houston":("TX",0.0210,1500,2600,1.20),
    "san antonio":("TX",0.0195,1200,2400,1.00),"austin":("TX",0.0190,1900,2400,1.55),
    "el paso":("TX",0.0180,1100,2200,0.85),"plano":("TX",0.0200,1700,2500,1.40),
    "arlington":("TX",0.0200,1400,2500,1.10),"oklahoma city":("OK",0.0110,1000,2400,0.85),
    "tulsa":("OK",0.0095,1000,2300,0.80),"albuquerque":("NM",0.0072,1100,1500,0.85),
    "tucson":("AZ",0.0062,1100,1300,0.90),"phoenix":("AZ",0.0060,1500,1300,1.20),
    "mesa":("AZ",0.0058,1400,1300,1.10),"scottsdale":("AZ",0.0060,1900,1400,1.55),
    "chandler":("AZ",0.0058,1500,1300,1.20),"tempe":("AZ",0.0060,1500,1300,1.20),
    "henderson":("NV",0.0055,1400,950,1.15),"las vegas":("NV",0.0055,1300,950,1.05),
    "reno":("NV",0.0055,1400,1000,1.10),"los angeles":("CA",0.0110,2600,1500,2.10),
    "la":("CA",0.0110,2600,1500,2.10),"san francisco":("CA",0.0075,3400,1600,2.80),
    "sf":("CA",0.0075,3400,1600,2.80),"san jose":("CA",0.0075,2900,1500,2.40),
    "san diego":("CA",0.0080,2400,1600,1.95),"sacramento":("CA",0.0080,1800,1300,1.45),
    "fresno":("CA",0.0070,1300,1200,1.00),"long beach":("CA",0.0110,2200,1500,1.80),
    "oakland":("CA",0.0075,2400,1500,1.95),"anaheim":("CA",0.0110,2200,1500,1.80),
    "bakersfield":("CA",0.0070,1300,1200,1.00),"seattle":("WA",0.0092,2200,1100,1.80),
    "spokane":("WA",0.0092,1200,1000,0.95),"tacoma":("WA",0.0092,1600,1000,1.30),
    "portland":("OR",0.0085,1700,900,1.35),"denver":("CO",0.0052,1900,1800,1.55),
    "aurora":("CO",0.0050,1600,1700,1.30),"colorado springs":("CO",0.0050,1500,1800,1.20),
    "fort collins":("CO",0.0053,1700,1800,1.40),"salt lake city":("UT",0.0053,1500,850,1.20),
    "provo":("UT",0.0052,1300,800,1.05),"boise":("ID",0.0063,1400,950,1.10),
    "anchorage":("AK",0.0120,1600,1000,1.30),"honolulu":("HI",0.0030,2400,600,2.00),
}

# Zip prefix (first 3 digits) → city key for common markets
_ZIP_PREFIX_CITY = {
    "100":"new york city","101":"new york city","102":"new york city","103":"new york city",
    "104":"new york city","110":"new york city","111":"new york city","112":"brooklyn",
    "113":"queens","114":"queens","021":"boston","022":"boston","024":"boston",
    "191":"philadelphia","192":"philadelphia","606":"chicago","607":"chicago","608":"chicago",
    "770":"houston","771":"houston","772":"houston","773":"houston","774":"houston",
    "750":"dallas","751":"dallas","752":"dallas","753":"dallas","760":"fort worth",
    "787":"austin","786":"austin","900":"los angeles","901":"los angeles","902":"los angeles",
    "903":"los angeles","904":"los angeles","905":"los angeles","910":"los angeles",
    "911":"los angeles","913":"long beach","940":"san francisco","941":"san francisco",
    "921":"san diego","922":"san diego","980":"seattle","981":"seattle","972":"portland",
    "973":"portland","802":"denver","800":"denver","801":"denver","331":"miami",
    "330":"miami","320":"jacksonville","328":"orlando","336":"tampa","303":"atlanta",
    "282":"charlotte","276":"raleigh","370":"nashville","481":"detroit","432":"columbus",
    "441":"cleveland","452":"cincinnati","530":"milwaukee","554":"minneapolis",
    "631":"st. louis","641":"kansas city",
}


def detect_location(address):
    """Return (label, state, tax_rate, city_rent, annual_insurance, rent_sqft) from address."""
    if not address:
        return None, None, None, None, None, None
    text  = address.strip()
    lower = text.lower()

    # 1. Try zip code first (most specific)
    zip_m = re.search(r'\b(\d{5})\b', text)
    if zip_m:
        city_key = _ZIP_PREFIX_CITY.get(zip_m.group(1)[:3])
        if city_key and city_key in _CITY_DATA:
            st, tax, rent, ins, rsqft = _CITY_DATA[city_key]
            return city_key.title() + " (zip " + zip_m.group(1) + ")", st, tax, rent, ins, rsqft

    # 2. Try city name match from address parts
    parts = [p.strip().lower() for p in re.split(r'[,\n]', lower)]
    for part in parts:
        clean = re.sub(r'\b[a-z]{2}\b$', '', part).strip()
        clean = re.sub(r'\b\d+\b', '', clean).strip().rstrip(',').strip()
        if clean in _CITY_DATA:
            st, tax, rent, ins, rsqft = _CITY_DATA[clean]
            return clean.title(), st, tax, rent, ins, rsqft

    # 3. Fall back to state
    state = None
    m = re.search(r'[,\s]([A-Z]{2})(?:\s+\d{5})?(?:\s*$|,)', text)
    if m and m.group(1) in _STATE_TAX_RATE:
        state = m.group(1)
    if not state:
        for name, code in _STATE_NAMES.items():
            if name in lower:
                state = code
                break
    if state:
        return state + " (state avg)", state, _STATE_TAX_RATE.get(state), None, _STATE_INSURANCE.get(state), _STATE_RENT_SQFT.get(state)

    return None, None, None, None, None, None


def estimate_location_costs(address, purchase_price, sqft=0):
    """Return (monthly_taxes, monthly_insurance, monthly_utilities, monthly_rent, label).
    Tries Census Geocoder → county-level data first, then city/state fallback."""
    # ── Try Census Geocoder for county-level precision ────────────────────────
    geoid, county_name, _ = fetch_census_county(address)
    if geoid and purchase_price:
        county_tax  = _COUNTY_TAX_RATE.get(geoid)
        county_ins  = _COUNTY_INSURANCE.get(geoid)
        # Still use detect_location for rent/utilities/state
        loc_label, state, _, city_rent, _, rent_sqft = detect_location(address)
        if county_tax:
            monthly_taxes     = round((purchase_price * county_tax) / 12)
            monthly_insurance = round(county_ins / 12) if county_ins else (
                round(_STATE_INSURANCE.get(state, 1200) / 12) if state else None
            )
            monthly_utilities = _STATE_UTILITIES.get(state) if state else None
            if sqft and sqft > 0 and rent_sqft:
                monthly_rent = round(sqft * rent_sqft)
            elif city_rent:
                monthly_rent = city_rent
            elif purchase_price and state and state in _STATE_RENT_YIELD:
                monthly_rent = round(purchase_price * _STATE_RENT_YIELD[state] / 12)
            else:
                monthly_rent = None
            label = county_name + " (county)"
            return monthly_taxes, monthly_insurance, monthly_utilities, monthly_rent, label

    # ── Fallback: city/zip/state lookup ──────────────────────────────────────
    label, state, tax_rate, city_rent, annual_ins, rent_sqft = detect_location(address)
    if not label:
        return None, None, None, None, None
    monthly_taxes     = round((purchase_price * tax_rate) / 12) if tax_rate and purchase_price else None
    monthly_insurance = round(annual_ins / 12) if annual_ins else None
    monthly_utilities = _STATE_UTILITIES.get(state) if state else None
    if sqft and sqft > 0 and rent_sqft:
        monthly_rent = round(sqft * rent_sqft)
    elif city_rent:
        monthly_rent = city_rent
    elif purchase_price and state in _STATE_RENT_YIELD:
        monthly_rent = round(purchase_price * _STATE_RENT_YIELD[state] / 12)
    else:
        monthly_rent = None
    return monthly_taxes, monthly_insurance, monthly_utilities, monthly_rent, label


def fetch_hud_fmr(address, bedrooms, token):
    """Fetch HUD Fair Market Rent by matching address city to metro area.
    Returns (rent, label, error)."""
    try:
        import requests as _req
    except ImportError:
        return None, None, "requests package required: pip install requests"

    # Extract state from address
    state = None
    m = re.search(r'[,\s]([A-Z]{2})(?:\s+\d{5})?(?:\s*(?:$|,))', address)
    if m and m.group(1) in _STATE_TAX_RATE:
        state = m.group(1)
    if not state:
        for name, code in _STATE_NAMES.items():
            if name in address.lower():
                state = code
                break
    if not state:
        return None, None, "No state found in address — include state abbreviation (e.g. TX)"

    headers = {"Authorization": f"Bearer {token}"}

    # Fetch all metro FMR data for the state
    try:
        resp = _req.get(
            f"https://www.huduser.gov/hudapi/public/fmr/statedata/{state}",
            headers=headers, timeout=15,
        )
        if resp.status_code == 401:
            return None, None, "HUD token invalid or expired"
        if resp.status_code != 200:
            return None, None, f"HUD API error: {resp.status_code}"
        metros = resp.json().get("data", {}).get("metroareas", [])
        if not metros:
            return None, None, f"No metro FMR data available for {state}"
    except Exception as e:
        return None, None, f"HUD API request failed: {e}"

    # Build city candidates from address
    addr_lower = address.lower()
    parts = [p.strip() for p in re.split(r'[,\n]', addr_lower)]
    city_candidates = []
    for part in parts:
        clean = re.sub(r'\b[a-z]{2}\b$', '', part).strip()  # strip trailing state abbrev
        clean = re.sub(r'\b\d+\b', '', clean).strip()        # strip numbers
        if clean and len(clean) > 2:
            city_candidates.append(clean)

    # Score each metro by how well city candidates match
    best_metro = None
    best_score = 0
    for metro in metros:
        metro_lower = metro["metro_name"].lower()
        for city in city_candidates:
            # Split city into words and check for word-level matches
            words = [w for w in city.split() if len(w) > 2]
            score = sum(len(w) for w in words if w in metro_lower)
            if score > best_score:
                best_score = score
                best_metro = metro

    if not best_metro or best_score < 3:
        sample = ", ".join(m["metro_name"] for m in metros[:4])
        return None, None, f"Could not match address to a {state} metro area. Sample metros: {sample}"

    br_map = {
        "Studio": "Efficiency",
        "1 BR":   "One-Bedroom",
        "2 BR":   "Two-Bedroom",
        "3 BR":   "Three-Bedroom",
        "4 BR":   "Four-Bedroom",
    }
    field = br_map.get(bedrooms, "Two-Bedroom")
    rent = best_metro.get(field)
    if rent is None:
        return None, None, f"No {bedrooms} FMR data for {best_metro['metro_name']}"

    return int(rent), f"HUD FMR — {best_metro['metro_name']}", None


# infer_units_from_type and populate_sidebar_from_data imported from dq_utils

# ── Sidebar — Manual Assumptions ─────────────────────────────────────────────

with st.sidebar:
    if st.session_state.get("authenticated") and _get_users():
        uname = st.session_state.get("username", "")
        st.caption(f"Signed in as **{uname}**")
        if st.button("Sign Out", key="_logout_btn"):
            st.session_state["authenticated"] = False
            st.session_state.pop("username", None)
            st.rerun()
        st.markdown("<hr style='margin:8px 0'>", unsafe_allow_html=True)

    st.header("Manual Assumptions")
    st.caption("Fill any fields to override or skip document upload entirely.")

    with st.expander("Property", expanded=True):
        s_address = st.text_input("Address / Name", placeholder="e.g. 123 Main St", key="sb_address")
        s_type    = st.text_input("Property Type", placeholder="e.g. Retail Strip", key="sb_type")
        s_sqft    = st.number_input("Square Feet", min_value=0, step=500, key="sb_sqft")
        s_price   = st.number_input("Purchase Price ($)", min_value=0, step=10000, key="sb_price")
        s_down    = st.number_input("Down Payment (%)", min_value=0.0, max_value=100.0, step=1.0, key="sb_down")
        s_closing = st.number_input("Closing Costs ($)", min_value=0, step=1000, key="sb_closing")

        # Auto-calculate loan
        _auto_loan = int(s_price * (1 - s_down / 100.0)) if s_price > 0 else 0
        st.caption("Loan Amount (auto): " + fmt_d(_auto_loan))
        if s_closing == 0 and _auto_loan > 0:
            st.caption("Closing Costs (default 2%): " + fmt_d(_auto_loan * 0.02))

    with st.expander("Income", expanded=True):
        _addr_r    = st.session_state.get("sb_address", "")
        _price_r   = st.session_state.get("sb_price", 0)
        _hud_token = st.session_state.get("sb_hud_token", "") or HUD_TOKEN

        s_units = st.number_input("Number of Units", min_value=1, step=1, key="sb_units")

        # Bedrooms selector (used by HUD FMR and estimate button)
        s_bedrooms = st.selectbox(
            "Bedrooms (for rent estimate)",
            ["Studio", "1 BR", "2 BR", "3 BR", "4 BR"],
            index=["Studio", "1 BR", "2 BR", "3 BR", "4 BR"].index(
                st.session_state.get("sb_bedrooms", "2 BR")
            ),
            key="sb_bedrooms",
        )

        if _addr_r and _price_r:
            _, _, _, _est_rent_loc, _rlabel = estimate_location_costs(
                _addr_r, _price_r, st.session_state.get("sb_sqft", 0)
            )
            if _hud_token:
                _hud_rent, _hud_label, _hud_err = fetch_hud_fmr(_addr_r, s_bedrooms, _hud_token)
                if _hud_rent:
                    _hud_total = _hud_rent * int(s_units)
                    if st.button(f"Estimate rent — {_hud_label}", use_container_width=True):
                        st.session_state["sb_gmi"] = _hud_total
                        st.rerun()
                    if s_units > 1:
                        st.caption(f"HUD FMR {s_bedrooms}: {fmt_d(_hud_rent)}/unit × {int(s_units)} units = {fmt_d(_hud_total)}/mo")
                    else:
                        st.caption(f"HUD FMR {s_bedrooms}: {fmt_d(_hud_rent)}/mo")
                else:
                    st.warning(f"HUD FMR failed: {_hud_err}", icon="⚠️")
                    if _rlabel and st.button(f"Estimate rent — {_rlabel} (local fallback)", use_container_width=True):
                        if _est_rent_loc:
                            st.session_state["sb_gmi"] = _est_rent_loc
                        st.rerun()
            elif _rlabel:
                if st.button(f"Estimate rent — {_rlabel}", use_container_width=True):
                    if _est_rent_loc:
                        st.session_state["sb_gmi"] = _est_rent_loc
                    st.rerun()
                _rent_basis = "sq ft" if st.session_state.get("sb_sqft", 0) > 0 else "purchase price"
                st.caption(f"~{fmt_d(_est_rent_loc)}/mo (based on {_rent_basis})")

        s_gmi     = st.number_input("Gross Monthly Income ($)", min_value=0, step=500, key="sb_gmi")
        s_vacancy = st.number_input("Vacancy Rate (%)", min_value=0.0, max_value=100.0, step=1.0, key="sb_vacancy")

    with st.expander("Operating Expenses", expanded=True):
        # Estimate button
        _addr  = st.session_state.get("sb_address", "")
        _price = st.session_state.get("sb_price", 0)
        if _addr and _price:
            _et, _ei, _eu, _, _elabel = estimate_location_costs(_addr, _price, st.session_state.get("sb_sqft", 0))
            if _elabel and st.button(f"Estimate taxes, insurance & utilities — {_elabel}", use_container_width=True):
                if _et: st.session_state["sb_taxes"]  = _et
                if _ei: st.session_state["sb_insure"] = _ei
                if _eu: st.session_state["sb_util"]   = _eu
                st.rerun()
            if _elabel:
                st.caption(f"~{fmt_d(_et)}/mo taxes · ~{fmt_d(_ei)}/mo insurance · ~{fmt_d(_eu)}/mo utilities")

        st.markdown("**Fixed ($/month)**")
        s_taxes   = st.number_input("Property Taxes ($/yr ÷ 12)", min_value=0, step=100, key="sb_taxes")
        s_insure  = st.number_input("Insurance ($/yr ÷ 12)", min_value=0, step=50, key="sb_insure")
        s_util    = st.number_input("Utilities ($/mo)", min_value=0, step=50, key="sb_util")
        st.markdown("**Variable (% of gross rent)**")
        s_capex   = st.number_input("CapEx (%)", min_value=0.0, max_value=50.0, step=0.5, key="sb_capex")
        s_maint   = st.number_input("Maintenance (%)", min_value=0.0, max_value=50.0, step=0.5, key="sb_maint")
        s_mgmt    = st.number_input("Management (%)", min_value=0.0, max_value=50.0, step=0.5, key="sb_mgmt")

        # Show auto-calculated total
        _var_mo  = s_gmi * (s_capex + s_maint + s_mgmt) / 100.0
        _fix_mo  = float(s_taxes) + float(s_insure) + float(s_util)
        _total_opex = _var_mo + _fix_mo
        st.caption("Total monthly expenses (auto): " + fmt_d(_total_opex))

    with st.expander("Financing"):
        s_rate  = st.number_input("Interest Rate (%)", min_value=0.0, max_value=30.0, step=0.25, key="sb_rate")
        s_amort = st.number_input("Amortization (years)", min_value=1, max_value=40, step=1, key="sb_amort")

        # Show auto-calculated mortgage
        if s_price > 0 and s_rate > 0:
            _loan_preview = s_price * (1 - s_down / 100.0)
            _pmt_preview  = calc_payment(_loan_preview, s_rate / 100.0, int(s_amort))
            st.caption("Est. monthly mortgage: " + fmt_d(_pmt_preview))

    with st.expander("Growth Assumptions"):
        s_appr   = st.number_input("Appreciation Rate (%)", min_value=0.0, max_value=20.0, step=0.25, key="sb_appr")
        s_rent_g = st.number_input("Rent Growth Rate (%)", min_value=0.0, max_value=20.0, step=0.25, key="sb_rent_g")

    manual_btn = st.button("Run Deal", type="primary")

    if manual_btn:
        price   = float(s_price)
        down_p  = s_down / 100.0
        loan    = price * (1 - down_p)
        gmi_f   = float(s_gmi)
        closing = float(s_closing) if s_closing > 0 else loan * 0.02

        # Build monthly opex from individual line items
        var_monthly  = gmi_f * (s_capex + s_maint + s_mgmt) / 100.0
        fix_monthly  = float(s_taxes) + float(s_insure) + float(s_util)
        total_monthly_opex = var_monthly + fix_monthly

        manual_data = {
            "property": {
                "address":                    s_address or None,
                "property_type":              s_type or None,
                "num_units":                  int(s_units),
                "square_feet":                float(s_sqft) if s_sqft > 0 else None,
                "purchase_price":             price if price > 0 else None,
                "down_payment_pct":           down_p,
                "gross_monthly_income":       gmi_f if gmi_f > 0 else None,
                "vacancy_rate":               s_vacancy / 100.0,
                "monthly_operating_expenses": total_monthly_opex if total_monthly_opex > 0 else None,
                "operating_expenses_pct":     0.45,
                "appreciation_rate":          s_appr / 100.0,
                "rent_growth_rate":           s_rent_g / 100.0,
                "annual_cap_ex":              None,
                "closing_costs":              closing,
                # line-item detail for breakdown section (prefixed with _)
                "_taxes":     float(s_taxes) if s_taxes > 0 else 0,
                "_insurance": float(s_insure) if s_insure > 0 else 0,
                "_utilities": float(s_util) if s_util > 0 else 0,
                "_capex_pct": s_capex / 100.0 if s_capex > 0 else 0,
                "_maint_pct": s_maint / 100.0 if s_maint > 0 else 0,
                "_mgmt_pct":  s_mgmt / 100.0 if s_mgmt > 0 else 0,
            },
            "financing": {
                "interest_rate":      s_rate / 100.0,
                "amortization_years": int(s_amort),
                "loan_amount":        loan if loan > 0 else None,
            },
            "missing_info": [],
        }
        st.session_state["dashboard_data"] = manual_data
        save_to_history(manual_data)
        st.rerun()


# ── Main UI ───────────────────────────────────────────────────────────────────

st.title("Deal Analyzer")
st.markdown("Paste a Zillow link, upload deal documents, or fill in the sidebar manually.")

# ── Zillow input ──────────────────────────────────────────────────────────────
zillow_col, zillow_btn_col = st.columns([5, 1])
with zillow_col:
    zillow_input = st.text_input("Zillow URL or Address", placeholder="https://www.zillow.com/homedetails/... or 123 Main St, City, ST", key="zillow_input")
with zillow_btn_col:
    st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
    zillow_btn = st.button("Fetch from Zillow", type="primary", disabled=not zillow_input)

if zillow_btn and zillow_input:
    st.session_state.pop("dashboard_data", None)
    client = anthropic.Anthropic(api_key=API_KEY)
    with st.spinner("Fetching Zillow listing..."):
        zillow_text, err = fetch_zillow_text(zillow_input)
    if err == "BLOCKED":
        st.info("Couldn't pull Zillow listing — address saved. Fill in purchase price and rent in the sidebar, then click **Run Deal**.")
        default_data = {
            "property": {
                "address": zillow_input,
                "property_type": None,
                "num_units": 1,
                "square_feet": None,
                "purchase_price": None,
                "down_payment_pct": 0.25,
                "gross_monthly_income": None,
                "vacancy_rate": 0.0,
                "monthly_operating_expenses": None,
                "operating_expenses_pct": 0.45,
                "appreciation_rate": 0.03,
                "rent_growth_rate": 0.02,
                "annual_cap_ex": None,
                "closing_costs": None,
            },
            "financing": {
                "interest_rate": 0.07,
                "amortization_years": 30,
                "loan_amount": None,
            },
            "missing_info": ["purchase_price", "gross_monthly_income"],
        }
        st.session_state["dashboard_data"] = default_data
        st.session_state.setdefault("_sb_pending", {})["sb_address"] = zillow_input
        st.rerun()
    elif err:
        st.error(err)
    else:
        with st.spinner("Analyzing property data..."):
            try:
                content = [{"type": "text", "text": zillow_text + "\n\nExtract all deal data from this Zillow listing and return the JSON as specified."}]
                full_text = ""
                placeholder = st.empty()
                with client.messages.stream(
                    model=MODEL,
                    max_tokens=4096,
                    system=SYSTEM_PROMPT,
                    messages=[{"role": "user", "content": content}],
                ) as stream:
                    for chunk in stream.text_stream:
                        full_text += chunk
                        placeholder.code(full_text[:800] + ("..." if len(full_text) > 800 else ""), language="json")
                placeholder.empty()
                data = extract_json(full_text)
                if data:
                    st.session_state["dashboard_data"] = data
                    save_to_history(data)
                    populate_sidebar_from_data(data)
                    st.rerun()
                else:
                    st.error("Could not extract deal data from this Zillow listing.")
            except Exception as e:
                st.error("API error: " + str(e))

# ── Paste listing text ────────────────────────────────────────────────────────
pasted_text = st.text_area(
    "Or paste listing text directly",
    height=160,
    placeholder="Copy all text from any listing page (Ctrl+A → Ctrl+C) and paste here — Zillow, LoopNet, CoStar, MLS, etc.",
    key="pasted_text",
)
paste_btn = st.button("Analyze Pasted Text", type="primary", disabled=not pasted_text)
if paste_btn and pasted_text:
    client = anthropic.Anthropic(api_key=API_KEY)
    with st.spinner("Analyzing..."):
        try:
            content = [{"type": "text", "text": pasted_text + "\n\nExtract all deal data from this listing and return the JSON as specified."}]
            full_text = ""
            placeholder = st.empty()
            with client.messages.stream(
                model=MODEL, max_tokens=4096, system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": content}],
            ) as stream:
                for chunk in stream.text_stream:
                    full_text += chunk
                    placeholder.code(full_text[:800] + ("..." if len(full_text) > 800 else ""), language="json")
            placeholder.empty()
            data = extract_json(full_text)
            if data:
                st.session_state["dashboard_data"] = data
                save_to_history(data)
                populate_sidebar_from_data(data)
                st.rerun()
            else:
                st.error("Could not extract deal data from the pasted text.")
        except Exception as e:
            st.error("API error: " + str(e))

st.divider()

# ── Document upload ───────────────────────────────────────────────────────────
uploaded_files = st.file_uploader(
    "Or upload deal documents",
    accept_multiple_files=True,
    type=["pdf", "xlsx", "xls", "csv", "png", "jpg", "jpeg", "docx", "doc"],
)

analyze_btn = st.button("Analyze with Documents", type="primary", disabled=not uploaded_files)

if analyze_btn and uploaded_files:
    st.session_state.pop("dashboard_data", None)

    client = anthropic.Anthropic(api_key=API_KEY)

    with st.spinner("Reading documents and analyzing deal..."):
        try:
            file_ids = upload_files(client, uploaded_files)

            if not file_ids:
                st.error("No files could be processed. Check that your document is a supported format.")
            else:
                content  = build_content(file_ids)

                full_text = ""
                placeholder = st.empty()

                with client.messages.stream(
                    model=MODEL,
                    max_tokens=4096,
                    system=SYSTEM_PROMPT,
                    messages=[{"role": "user", "content": content}],
                    extra_headers={"anthropic-beta": "files-api-2025-04-14"},
                ) as stream:
                    for chunk in stream.text_stream:
                        full_text += chunk
                        placeholder.code(full_text[:800] + ("..." if len(full_text) > 800 else ""), language="json")

                placeholder.empty()

                data = extract_json(full_text)
                if data:
                    st.session_state["dashboard_data"] = data
                    save_to_history(data)
                    populate_sidebar_from_data(data)
                    st.rerun()
                else:
                    st.error("Could not extract deal data from this document.")
                    with st.expander("See Claude's raw response (for debugging)"):
                        st.text(full_text[:3000] if full_text else "(empty response)")
                    st.info("Tip: The document may not contain financial data Claude can recognize. Try a rent roll, lease, or PSA — or enter assumptions manually in the sidebar.")

        except Exception as e:
            st.error("API error: " + str(e))

if "dashboard_data" in st.session_state:
    show_dashboard(st.session_state["dashboard_data"])
